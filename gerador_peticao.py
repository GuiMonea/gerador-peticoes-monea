
from fastapi import FastAPI, Request
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

app = FastAPI()

@app.post("/gerar-peticao")
async def gerar_peticao(request: Request):
    data = await request.json()

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Montserrat'
    font.size = Pt(12)

    def par(text, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, bold=False, size=12, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Inches(0.79)
        p.alignment = align
        return p

    # Cabeçalho
    par(f"AO JUÍZO DE DIREITO DA __ª VARA CÍVEL DA COMARCA DE {data['comarca']} – {data['estado']}.")
    doc.add_paragraph("")

    # Número do processo
    par(f"PROCESSO Nº: {data['numero_processo']}")

    # Qualificação
    par(f"{data['nome_autor']}, já qualificada nos autos do processo epigrafado, vem, respeitosamente, à presença de Vossa Excelência, por meio de seus procuradores infra-assinados, apresentar os presentes")

    # Título da peça
    p = doc.add_paragraph(f"{data['titulo_peticao']}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(247, 150, 70)

    # Fatos
    par("I. DOS FATOS", bold=True)
    par(data['fatos'])

    # Direito
    par("II. DO DIREITO", bold=True)
    par(data['fundamento'])

    # Pedidos
    par("III. DOS PEDIDOS", bold=True)
    par(data['pedidos'])

    # Final
    par("Por fim, que todas as intimações sejam realizadas exclusivamente em nome do advogado FELIPE AUGUSTO NUNES MONEA, com endereço profissional à Rua Fidêncio Ramos, 160 – CJ 706 e 707 – Vila Olímpia, São Paulo – SP, 04551-010 e endereço eletrônico publicacoes@monea.adv.br, sob pena de nulidade.")
    par("Termos em que,")
    par("Pede deferimento.")
    par(f"São Paulo/SP, {datetime.now().strftime('%d de %B de %Y')}.")

    par("FELIPE AUGUSTO NUNES MONEA\nOAB/SP 397.029")

    output_path = "/tmp/peticao_gerada.docx"
    doc.save(output_path)
    return FileResponse(output_path, filename="peticao_gerada.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
